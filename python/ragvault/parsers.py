"""Pluggable document parsers.

Built-in (no extra dependencies): txt, markdown, html, json, jsonl, csv and
source code. PDF and DOCX are optional extras with actionable error
messages. Parsers return structure (title, format) rather than flattening
everything blindly; failures raise :class:`IngestionError` with file,
parser, stage and a suggested fix.
"""

from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Callable, Optional

from .errors import IngestionError

CODE_EXTENSIONS = {
    ".py", ".rs", ".js", ".ts", ".tsx", ".jsx", ".go", ".java", ".c", ".h",
    ".cpp", ".hpp", ".rb", ".php", ".swift", ".kt", ".scala", ".sh", ".sql",
    ".yaml", ".yml", ".toml", ".ini", ".cfg",
}


@dataclass
class ParsedDocument:
    text: str
    format: str
    title: Optional[str] = None
    metadata: dict = field(default_factory=dict)


class _TextExtractor(HTMLParser):
    _SKIP = {"script", "style", "noscript", "template"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.title_parts: list[str] = []
        self._skip_depth = 0
        self._in_title = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._SKIP:
            self._skip_depth += 1
        if tag == "title":
            self._in_title = True
        if tag in {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP and self._skip_depth > 0:
            self._skip_depth -= 1
        if tag == "title":
            self._in_title = False

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if self._in_title:
            self.title_parts.append(data)
        else:
            self.parts.append(data)


def _parse_text(path: Path, raw: bytes) -> ParsedDocument:
    text = raw.decode("utf-8", errors="replace")
    return ParsedDocument(text=text, format="text")


def _parse_markdown(path: Path, raw: bytes) -> ParsedDocument:
    text = raw.decode("utf-8", errors="replace")
    title = None
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            break
    return ParsedDocument(text=text, format="markdown", title=title)


def _parse_html(path: Path, raw: bytes) -> ParsedDocument:
    extractor = _TextExtractor()
    extractor.feed(raw.decode("utf-8", errors="replace"))
    text = "\n".join(
        line.strip() for line in "".join(extractor.parts).splitlines() if line.strip()
    )
    title = "".join(extractor.title_parts).strip() or None
    return ParsedDocument(text=text, format="html", title=title)


def _parse_json(path: Path, raw: bytes) -> ParsedDocument:
    try:
        data = json.loads(raw.decode("utf-8", errors="replace"))
    except json.JSONDecodeError as exc:
        raise IngestionError(
            f"invalid JSON: {exc}", file=str(path), parser="json", stage="decode",
            suggestion="validate the file with `python -m json.tool`",
        ) from exc
    text = json.dumps(data, ensure_ascii=False, indent=2)
    return ParsedDocument(text=text, format="json")


def _parse_jsonl(path: Path, raw: bytes) -> ParsedDocument:
    lines = []
    for i, line in enumerate(raw.decode("utf-8", errors="replace").splitlines()):
        if not line.strip():
            continue
        try:
            data = json.loads(line)
        except json.JSONDecodeError as exc:
            raise IngestionError(
                f"invalid JSONL at line {i + 1}: {exc}", file=str(path),
                parser="jsonl", stage="decode",
            ) from exc
        if isinstance(data, dict) and "text" in data:
            lines.append(str(data["text"]))
        else:
            lines.append(json.dumps(data, ensure_ascii=False))
    return ParsedDocument(text="\n\n".join(lines), format="jsonl")


def _parse_csv(path: Path, raw: bytes) -> ParsedDocument:
    text = raw.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ParsedDocument(text="", format="csv")
    header = rows[0]
    lines = []
    for row in rows[1:]:
        pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
        if pairs:
            lines.append("; ".join(pairs))
    return ParsedDocument(text="\n\n".join(lines), format="csv", metadata={"columns": header})


def _parse_code(path: Path, raw: bytes) -> ParsedDocument:
    text = raw.decode("utf-8", errors="replace")
    return ParsedDocument(
        text=text, format="code", metadata={"language": path.suffix.lstrip(".")}
    )


def _parse_pdf(path: Path, raw: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise IngestionError(
            "PDF support requires an optional dependency", file=str(path),
            parser="pdf", stage="import",
            suggestion='pip install "ragvault[pdf]"',
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages = [(i + 1, page.extract_text() or "") for i, page in enumerate(reader.pages)]
    except Exception as exc:  # pypdf raises many exception types
        raise IngestionError(
            f"failed to read PDF: {exc}", file=str(path), parser="pdf", stage="extract",
            suggestion="check the file opens in a PDF viewer; encrypted PDFs are unsupported",
        ) from exc
    text = "\n\n".join(f"[page {n}]\n{t}" for n, t in pages if t.strip())
    meta = reader.metadata
    title = (meta.title if meta else None) or None
    return ParsedDocument(text=text, format="pdf", title=title,
                          metadata={"pages": len(pages)})


def _parse_docx(path: Path, raw: bytes) -> ParsedDocument:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:
        raise IngestionError(
            "DOCX support requires an optional dependency", file=str(path),
            parser="docx", stage="import",
            suggestion='pip install "ragvault[office]"',
        ) from exc
    try:
        document = docx.Document(io.BytesIO(raw))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    except Exception as exc:
        raise IngestionError(
            f"failed to read DOCX: {exc}", file=str(path), parser="docx", stage="extract",
        ) from exc
    return ParsedDocument(text="\n\n".join(paragraphs), format="docx")


_PARSERS: dict[str, Callable[[Path, bytes], ParsedDocument]] = {
    ".txt": _parse_text,
    ".text": _parse_text,
    ".md": _parse_markdown,
    ".markdown": _parse_markdown,
    ".html": _parse_html,
    ".htm": _parse_html,
    ".json": _parse_json,
    ".jsonl": _parse_jsonl,
    ".ndjson": _parse_jsonl,
    ".csv": _parse_csv,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
}


def supported_extensions() -> set[str]:
    return set(_PARSERS) | CODE_EXTENSIONS


def register_parser(extension: str, parser: Callable[[Path, bytes], ParsedDocument]) -> None:
    """Register a custom parser for a file extension (plugin point)."""
    _PARSERS[extension.lower()] = parser


def parse_file(path: str | Path) -> ParsedDocument:
    path = Path(path)
    suffix = path.suffix.lower()
    parser = _PARSERS.get(suffix)
    if parser is None and suffix in CODE_EXTENSIONS:
        parser = _parse_code
    if parser is None:
        parser = _parse_text  # unknown extensions degrade to plain text
    try:
        raw = path.read_bytes()
    except OSError as exc:
        raise IngestionError(
            f"cannot read file: {exc}", file=str(path), parser=suffix or "text",
            stage="read",
        ) from exc
    doc = parser(path, raw)
    if doc.title is None:
        doc.title = path.stem
    return doc
